"""
Servicio para extraer texto de diferentes tipos de documentos:
- PDF: utilizando pdfplumber
- DOCX: utilizando python-docx
- DOC: debe ser convertido primero (no soportado directamente)
- RTF: utilizando striprtf
"""
import io
import logging
import re
from typing import Optional

import pdfplumber
from docx import Document
from striprtf.striprtf import rtf_to_text

logger = logging.getLogger(__name__)


def sanitize_text(text: str) -> str:
    """
    Sanitiza el texto para asegurar que sea compatible con UTF-8 y PostgreSQL.

    Args:
        text: Texto a sanitizar

    Returns:
        Texto sanitizado
    """
    if not text:
        return ""

    # Eliminar bytes nulos (causantes del error en PostgreSQL)
    text = text.replace('\x00', '')

    # Eliminar caracteres de control excepto saltos de línea y tabs
    text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)

    # Intentar codificar y decodificar para asegurar que sea UTF-8 válido
    try:
        # Usar 'replace' para sustituir caracteres inválidos por '�'
        text = text.encode('utf-8', 'replace').decode('utf-8')
    except Exception as e:
        logger.warning(f"Error al sanitizar texto: {str(e)}")
        # En caso extremo, usar sólo ASCII
        text = ''.join(c if ord(c) < 128 else '?' for c in text)

    return text


def detect_binary_content(text: str, threshold: float = 0.3) -> bool:
    """
    Detecta si el texto parece ser contenido binario en lugar de texto real.

    Args:
        text: Texto a analizar
        threshold: Umbral de proporción de caracteres no imprimibles

    Returns:
        True si parece ser contenido binario
    """
    if not text:
        return False

    # Contar caracteres no imprimibles
    non_printable = 0
    for char in text:
        if char < ' ' and char != '\n' and char != '\t' and char != '\r':
            non_printable += 1

    # Si más de threshold% son no imprimibles, considerar binario
    return (non_printable / len(text)) > threshold


def extract_text_from_bytes(file_bytes: bytes, content_type: str) -> Optional[str]:
    """
    Extrae texto plano de archivos en diferentes formatos.
    
    Args:
        file_bytes: Bytes del archivo
        content_type: Tipo MIME del archivo (ej. 'application/pdf')
    
    Returns:
        String con el texto extraído o None si hubo un error
    """
    try:
        extracted_text = None

        if content_type == "application/pdf":
            extracted_text = extract_from_pdf(file_bytes)
        elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            extracted_text = extract_from_docx(file_bytes)
        elif content_type in ["application/rtf", "text/rtf"]:
            extracted_text = extract_from_rtf(file_bytes)
        elif content_type == "application/msword":
            logger.warning("Extracción directa de archivos .doc no soportada")
            return "El formato DOC requiere conversión previa"
        else:
            logger.error(f"Tipo de documento no soportado: {content_type}")
            return None

        # Si no hay texto o parece contenido binario, retornar mensaje apropiado
        if not extracted_text:
            logger.warning("No se pudo extraer texto del documento")
            return "No se pudo extraer texto del documento"

        if detect_binary_content(extracted_text):
            logger.warning("El contenido extraído parece ser binario, no texto")
            return "El documento contiene datos binarios que no pueden ser extraídos como texto"

        # Sanitizar el texto para asegurar compatibilidad con UTF-8 y PostgreSQL
        clean_text = sanitize_text(extracted_text)
        return clean_text

    except Exception as e:
        logger.error(f"Error extrayendo texto: {str(e)}")
        return f"Error al procesar el documento: {str(e)}"


def extract_from_pdf(pdf_bytes: bytes) -> str:
    """Extrae texto de un archivo PDF"""
    text_content = []
    
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            text_content.append(text)
    
    return "\n\n".join(text_content)


def extract_from_docx(docx_bytes: bytes) -> str:
    """Extrae texto de un archivo DOCX"""
    doc = Document(io.BytesIO(docx_bytes))
    
    # Extrae el texto de cada párrafo
    text_content = []
    for paragraph in doc.paragraphs:
        if paragraph.text:
            text_content.append(paragraph.text)
    
    # También extrae texto de tablas
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text:
                    row_text.append(cell.text.strip())
            if row_text:
                text_content.append(" | ".join(row_text))
    
    return "\n\n".join(text_content)


def extract_from_rtf(rtf_bytes: bytes) -> str:
    """Extrae texto de un archivo RTF"""
    # Examinar el inicio del archivo para diagnóstico
    try:
        sample = rtf_bytes[:100].decode('latin1', errors='replace')
        logger.debug(f"Primeros 100 bytes del archivo: {repr(sample)}")
        
        # Verificar diferentes posibles inicios de RTF
        if '{\\rtf' in sample:
            logger.debug("Archivo RTF estándar detectado")
        elif '\\rtf' in sample:
            logger.debug("Posible RTF con formato alternativo detectado")
        elif 'PK' in sample[:2]:
            logger.debug("El archivo parece ser un ZIP/DOCX/XLSX, no un RTF")
            return "El archivo parece ser un archivo comprimido (ZIP/DOCX), no un RTF"
        elif '%PDF' in sample:
            logger.debug("El archivo parece ser un PDF, no un RTF")
            return "El archivo parece ser un PDF, no un RTF"
        else:
            logger.debug("No se reconoce el formato del archivo")
    except Exception as e:
        logger.error(f"Error al examinar el inicio del archivo: {str(e)}")
    
    # Intentar múltiples métodos de extracción
    result_text = ""
    errors = []
    
    # 1. Método estándar - usando striprtf
    try:
        # Convertir bytes a string probando diferentes codificaciones
        encodings = ['utf-8', 'latin1', 'cp1252', 'iso-8859-1']
        rtf_text = None
        
        for encoding in encodings:
            try:
                rtf_text = rtf_bytes.decode(encoding, errors='replace')
                logger.debug(f"Decodificación exitosa con {encoding}")
                break
            except UnicodeDecodeError:
                continue
        
        if rtf_text:
            try:
                plain_text = rtf_to_text(rtf_text)
                if plain_text and len(plain_text) > 20:  # Verificar que el resultado sea significativo
                    return sanitize_text(plain_text)
            except Exception as e:
                errors.append(f"Error con striprtf: {str(e)}")
                logger.warning(f"Error al procesar RTF con striprtf: {str(e)}")
    except Exception as e:
        errors.append(f"Error general en método 1: {str(e)}")
    
    # 2. Método alternativo - extracción de texto simple
    try:
        if rtf_text:
            # Eliminar comandos RTF
            simple_text = re.sub(r'\\[a-zA-Z0-9]+\s?', ' ', rtf_text)
            # Eliminar llaves y otros caracteres de control
            simple_text = re.sub(r'[{}]|\\\n|\\\r', ' ', simple_text)
            # Eliminar múltiples espacios
            simple_text = re.sub(r'\s+', ' ', simple_text).strip()
            
            if simple_text and len(simple_text) > 50:  # Resultado significativo
                return sanitize_text(simple_text)
    except Exception as e:
        errors.append(f"Error en método 2: {str(e)}")
    
    # 3. Último recurso - forzar decodificación e interpretación
    try:
        # Intentar extraer cualquier texto legible del archivo
        raw_text = rtf_bytes.decode('latin1', errors='replace')
        
        # Buscar bloques de texto legibles (secuencias de al menos 5 caracteres imprimibles)
        text_blocks = re.findall(r'[A-Za-z0-9áéíóúüñÁÉÍÓÚÜÑ.,;:¿?¡! ]{5,}', raw_text)
        
        if text_blocks:
            result = " ... ".join(text_blocks)
            return sanitize_text(result)
    except Exception as e:
        errors.append(f"Error en método 3: {str(e)}")
    
    # Si llegamos aquí, ningún método funcionó
    error_detail = "; ".join(errors)
    logger.error(f"No se pudo extraer texto del RTF: {error_detail}")
    return f"No se pudo extraer texto del archivo. Formato no compatible o archivo dañado."
